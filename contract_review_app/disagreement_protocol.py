"""Utility to generate disagreement protocol reports.

This module provides :func:`generate_disagreement_protocol` which creates
both DOCX and JSON reports summarising issues discovered during contract
analysis.  Each entry in the report includes the clause identifier,
identified issue, legal reference, recommendation and resolution status.

The function accepts the original analysis results and (optionally) the
results after fixes were applied.  When updated results are supplied the
function marks clauses as ``fixed`` when their status becomes ``OK`` or the
text changes; otherwise the clause is marked as ``pending``.
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List, Optional
import json

from docx import Document

from contract_review_app.core.schemas import AnalysisOutput


def _index_by_clause(outputs: Iterable[AnalysisOutput]) -> Dict[str, AnalysisOutput]:
    """Create a mapping from clause identifier to AnalysisOutput."""
    mapping: Dict[str, AnalysisOutput] = {}
    for out in outputs:
        key = out.clause_id or out.clause_type
        mapping[key] = out
    return mapping


def generate_disagreement_protocol(
    original: List[AnalysisOutput],
    updated: Optional[List[AnalysisOutput]] = None,
    docx_path: Path | str = "disagreement_protocol.docx",
    json_path: Path | str = "disagreement_protocol.json",
) -> None:
    """Generate disagreement protocol in DOCX and JSON formats.

    Parameters
    ----------
    original:
        List of :class:`AnalysisOutput` produced by the analysis engine.
    updated:
        Optional list of :class:`AnalysisOutput` after fixes were applied.
    docx_path:
        Path where the DOCX report will be written.
    json_path:
        Path where the JSON representation will be written.
    """

    updated_map = _index_by_clause(updated or [])
    items: List[Dict[str, str]] = []

    for clause in original:
        clause_id = clause.clause_id or clause.clause_type
        finding = clause.findings[0] if clause.findings else None
        issue = finding.message if finding else ""
        reference = (
            "; ".join(finding.legal_basis) if finding and finding.legal_basis else ""
        )
        recommendation = (
            clause.recommendations[0]
            if clause.recommendations
            else clause.proposed_text
        )

        status = "pending"
        if updated_map:
            u = updated_map.get(clause_id)
            if u and (u.status == "OK" or u.text != clause.text):
                status = "fixed"

        items.append(
            {
                "clauseId": str(clause_id),
                "issue": issue,
                "recommendation": recommendation,
                "reference": reference,
                "status": status,
            }
        )

    # Write JSON output
    json_path = Path(json_path)
    json_path.write_text(
        json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Write DOCX output
    doc = Document()
    doc.add_heading("Протокол разногласий", level=1)
    for item in items:
        doc.add_heading(item["clauseId"], level=2)
        p = doc.add_paragraph()
        p.add_run("Проблема: ").bold = True
        p.add_run(item["issue"] or "—")
        p = doc.add_paragraph()
        p.add_run("Рекомендация: ").bold = True
        p.add_run(item["recommendation"] or "—")
        p = doc.add_paragraph()
        p.add_run("Основание: ").bold = True
        p.add_run(item["reference"] or "—")
        p = doc.add_paragraph()
        p.add_run("Статус: ").bold = True
        p.add_run(item["status"])
    doc.save(docx_path)
