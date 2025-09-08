import tempfile
from pathlib import Path

import pytest
from docx import Document
from fpdf import FPDF

from contract_review_app.analysis.parser import parse_docx, parse_pdf


def test_parse_docx(tmp_path: Path):
    doc = Document()
    doc.add_heading("HEADING ONE", level=1)
    for i in range(1, 7):
        doc.add_paragraph(f"{i}. Item {i}")
    doc.add_paragraph("Regular paragraph.")
    file_path = tmp_path / "sample.docx"
    doc.save(file_path)
    data = file_path.read_bytes()
    parsed = parse_docx(data)
    assert len(parsed.normalized_text) > 0
    assert len(parsed.segments) >= 8
    for seg in parsed.segments:
        assert 0 <= seg["start"] < seg["end"] <= len(parsed.normalized_text)
    assert any(seg["kind"] == "heading" and seg.get("number") for seg in parsed.segments)


def test_parse_docx_multiline(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Line1\nLine2")
    file_path = tmp_path / "multi.docx"
    doc.save(file_path)
    parsed = parse_docx(file_path.read_bytes())
    assert "Line1\nLine2" in parsed.normalized_text


def test_parse_docx_large(tmp_path: Path) -> None:
    doc = Document()
    for i in range(500):
        doc.add_paragraph(f"Paragraph {i}")
    file_path = tmp_path / "large.docx"
    doc.save(file_path)
    parsed = parse_docx(file_path.read_bytes())
    assert len(parsed.normalized_text.splitlines()) >= 500


def test_parse_pdf(tmp_path: Path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, "First line\nSecond line\nThird line")
    file_path = tmp_path / "sample.pdf"
    pdf.output(str(file_path))
    data = file_path.read_bytes()
    try:
        parsed = parse_pdf(data)
    except NotImplementedError:
        pytest.skip("pdfminer not installed")
    assert len(parsed.normalized_text) > 0
    assert len(parsed.segments) >= 3
    for seg in parsed.segments:
        assert 0 <= seg["start"] < seg["end"] <= len(parsed.normalized_text)
