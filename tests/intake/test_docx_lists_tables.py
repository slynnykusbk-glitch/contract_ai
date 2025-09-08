from pathlib import Path

from docx import Document

from contract_review_app.utils.doc_loader import load_docx_text


def test_bulleted_and_numbered_lists(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Item 1", style="List Bullet")
    doc.add_paragraph("Item 2", style="List Bullet")
    doc.add_paragraph("Item 3", style="List Number")
    doc.add_paragraph("Item 4", style="List Number")
    file_path = tmp_path / "lists.docx"
    doc.save(file_path)

    text = load_docx_text(file_path)
    assert text.splitlines() == ["Item 1", "Item 2", "Item 3", "Item 4"]


def test_tables_merge_cells_and_newlines(tmp_path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    c = table.cell(1, 0)
    c.text = "A2 line1"
    c.add_paragraph("A2 line2")
    table.cell(1, 1).text = "B2"
    file_path = tmp_path / "table.docx"
    doc.save(file_path)

    text = load_docx_text(file_path)
    assert text == "A1\tB1\nA2 line1\nA2 line2\tB2"
