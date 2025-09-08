from pathlib import Path
from pathlib import Path

import pytest
from docx import Document

from contract_review_app.intake.parser import ParsedDocument

FIXTURES = [
    "tests/fixtures/intake_simple.txt",
    "tests/fixtures/intake_mixed.txt",
    "tests/fixtures/intake_fancy.txt",
]


@pytest.mark.parametrize("fname", FIXTURES)
def test_span_invariants(fname: str) -> None:
    raw = Path(fname).read_text(encoding="utf-8")
    doc = ParsedDocument.from_text(raw)

    assert all(
        0 <= doc.offset_map[i] < len(doc.content)
        and doc.offset_map[i] < doc.offset_map[i + 1]
        for i in range(len(doc.offset_map) - 1)
    )

    spans = [(int(s["start"]), int(s["end"])) for s in doc.segments]
    spans.sort()
    prev_end = 0
    for start, end in spans:
        assert prev_end <= start
        assert start < end
        prev_end = end

    nt = doc.normalized_text
    total = sum(1 for ch in nt if not ch.isspace())
    covered = 0
    for start, end in spans:
        covered += sum(1 for ch in nt[start:end] if not ch.isspace())
    if total:
        assert covered / total >= 0.995


def test_span_roundtrip_large_doc(tmp_path: Path) -> None:
    docx_doc = Document()
    for i in range(500):
        docx_doc.add_paragraph(f"Paragraph {i}")
    file_path = tmp_path / "large.docx"
    docx_doc.save(file_path)

    text = "\n".join(p.text for p in docx_doc.paragraphs)
    pd = ParsedDocument.from_text(text)

    span = (0, len(text))
    a, b = pd.map_raw_span_to_norm(*span)
    assert pd.map_norm_span_to_raw(a, b) == span

    word = "Paragraph 499"
    start = text.index(word)
    end = start + len(word)
    ns, ne = pd.map_raw_span_to_norm(start, end)
    assert pd.normalized_text[ns:ne] == word
    assert pd.map_norm_span_to_raw(ns, ne) == (start, end)
